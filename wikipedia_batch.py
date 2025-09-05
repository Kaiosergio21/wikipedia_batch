import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import quote
import re
import math

# Recebe o título do artigo
article = input("Digite o título do artigo da Wikipedia: ").strip()
safe_article = quote(article)
url = f"https://pt.wikipedia.org/wiki/{safe_article}"

# User-Agent
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
}

res = requests.get(url, headers=headers)

if res.status_code != 200:
    print(f"Erro ao acessar o artigo: {article}")
    exit()

soup = BeautifulSoup(res.text, "html.parser")

# Seleciona conteúdo principal
content = soup.find("div", {"id": "mw-content-text"})
if not content:
    print("Conteúdo do artigo não encontrado!")
    exit()

paragraphs = [p.get_text().strip() for p in content.find_all("p") if p.get_text().strip()]

if not paragraphs:
    print("Nenhum parágrafo encontrado no artigo!")
    exit()

# Pergunta a porcentagem que o usuário quer salvar
while True:
    try:
        perc = float(input(f"O artigo tem {len(paragraphs)} parágrafos. Qual porcentagem deseja salvar? (0-100): "))
        if 0 < perc <= 100:
            break
        else:
            print("Digite um valor entre 0 e 100.")
    except ValueError:
        print("Digite um número válido.")

# Pergunta a direção
while True:
    direction = input("Deseja pegar o texto de cima para baixo (C) ou de baixo para cima (B)? ").strip().upper()
    if direction in ["C", "B"]:
        break
    else:
        print("Digite 'C' para cima para baixo ou 'B' para baixo para cima.")

# Calcula quantos parágrafos salvar
num_to_save = math.ceil(len(paragraphs) * (perc / 100))

# Escolhe os parágrafos com base na direção
if direction == "C":
    selected_paragraphs = paragraphs[:num_to_save]
else:
    selected_paragraphs = paragraphs[-num_to_save:]

# Cria documento Word
doc = Document()
doc.add_heading(article, level=0)

for p in selected_paragraphs:
    p = re.sub(r'\[\d+\]', '', p)  # remove referências
    doc.add_paragraph(p)

filename = f"{article.replace(' ', '_')}_{perc}pct_{direction}.docx"
doc.save(filename)
print(f"{num_to_save} parágrafos salvos em {filename}")
